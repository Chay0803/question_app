# backend.py
import uvicorn
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
import tempfile, io, os, re, json, base64, hashlib
from typing import List, Tuple, Dict
import fitz  # PyMuPDF
import requests
import pandas as pd
from docx import Document
from io import BytesIO
from moviepy.editor import VideoFileClip
import subprocess
import whisper
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import datetime
from fastapi.encoders import jsonable_encoder
from fastapi import Depends, HTTPException, Security
from fastapi.security.api_key import APIKeyHeader
import httpx

MYSQL_USER = "root"
MYSQL_PASSWORD = "icfai123"
MYSQL_HOST = "localhost"
MYSQL_PORT = 3306
MYSQL_DB = "mcq_db"

DATABASE_URL = f"mysql+mysqlconnector://{MYSQL_USER}:{MYSQL_PASSWORD}@{MYSQL_HOST}:{MYSQL_PORT}/{MYSQL_DB}"

engine = create_engine(DATABASE_URL, pool_pre_ping=True)
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base = declarative_base()

class Question(Base):
    __tablename__ = "questions"

    id = Column(Integer, primary_key=True, index=True)
    topic = Column(String(255))
    type = Column(String(20))  # MCQ / Descriptive
    question = Column(Text, nullable=False)
    option_a = Column(Text)
    option_b = Column(Text)
    option_c = Column(Text)
    option_d = Column(Text)
    answer = Column(Text)
    descriptive_answer = Column(Text)
    difficulty = Column(String(10))
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

# Create table if not exists
Base.metadata.create_all(bind=engine)

# ---------- CONFIG ----------
# OLLAMA_URL = "http://localhost:11434/api/generate"  # change if required
# MODEL = "llama3"
HOST = "127.0.0.1"
PORT = 8000
API_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "x-ai/grok-4-fast:free"
API_KEY = "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"

INTERNAL_API_KEY = "myportalkey123"    
API_KEY_NAME = "X-API-Key"
api_key_header = APIKeyHeader(name=API_KEY_NAME, auto_error=False)

async def get_api_key(api_key_header: str = Security(api_key_header)):
    if api_key_header == INTERNAL_API_KEY:
        return api_key_header
    else:
        raise HTTPException(status_code=403, detail="Unauthorized: Invalid API Key")

# ---------- FASTAPI ----------
app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"], allow_credentials=True)

# Serve static files (put design.html and any assets inside ./static/)
static_dir = os.path.join(os.path.dirname(__file__), "static")
if not os.path.isdir(static_dir):
    os.makedirs(static_dir, exist_ok=True)
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# Serve design.html at root
@app.get("/", response_class=HTMLResponse)
async def index():
    fpath = os.path.join(static_dir, "design.html")
    if os.path.exists(fpath):
        return HTMLResponse(open(fpath, "r", encoding="utf-8").read())
    return HTMLResponse("<h3>Place design.html inside ./static/ and reload.</h3>")

# ---------- IN-MEMORY STATE & STORE ----------
IN_MEMORY_STORE = {}  # key -> {"data": bytes, "name": str, "mime": str}
STATE = {
    "pdf_uploads": 0,
    "last_pdf_hash": None,
    "last_pdf_pages": 0,
    "mcq_count": 0,
    "desc_count": 0
}

def store_result_bytes(key: str, data: bytes, filename: str, mime: str):
    IN_MEMORY_STORE[key] = {"data": data, "name": filename, "mime": mime}

@app.get("/download/{key}")
async def download_key(key: str):
    item = IN_MEMORY_STORE.get(key)
    if not item:
        return JSONResponse({"error": "Not found"}, status_code=404)
    return StreamingResponse(io.BytesIO(item["data"]), media_type=item["mime"],
                             headers={"Content-Disposition": f"attachment; filename={item['name']}"})

@app.get("/status")
async def status():
    """Return counters for the top dashboard (PDF uploads, pages, counts)."""
    return {
        "pdf_uploads": STATE["pdf_uploads"],
        "last_pdf_pages": STATE["last_pdf_pages"],
        "mcq_count": STATE["mcq_count"],
        "desc_count": STATE["desc_count"]
    }

# ---------- UTIL HELPERS (ported from your Streamlit code) ----------
def clean_text(text: str) -> str:
    if text is None:
        return ""
    return re.sub(r"[\x00-\x1F\x7F]", "", str(text))

def detect_index_range(doc, min_section_hits: int = 3, consecutive_break: int = 2) -> Tuple[int, int]:
    scores = []
    has_contents_flags = []
    for pno in range(doc.page_count):
        try:
            text = doc.load_page(pno).get_text("text") or ""
        except Exception:
            text = ""
        low = text.lower()
        has_contents = bool(re.search(r"\btable of contents\b|\bcontents\b", low))
        count_sections = len(re.findall(r"\b\d{1,2}\.\d+\b", text))
        count_leaders = len(re.findall(r"\.{2,}\s*\d+|\s+\d{1,3}\s*$", text, re.M))
        score = count_sections + 0.6 * count_leaders + (5 if has_contents else 0)
        scores.append(score)
        has_contents_flags.append(has_contents)

    if any(has_contents_flags):
        start_idx = next(i for i, f in enumerate(has_contents_flags) if f)
        end_idx = start_idx
        break_count = 0
        for i in range(start_idx + 1, len(scores)):
            if scores[i] >= 1.0:
                end_idx = i
                break_count = 0
            else:
                break_count += 1
                if break_count >= consecutive_break:
                    break
        return (start_idx + 1, end_idx + 1)

    start_idx = None
    for i, s in enumerate(scores):
        if s >= min_section_hits:
            start_idx = i
            break
    if start_idx is None:
        raise ValueError("Could not auto-detect contents/index pages.")

    end_idx = start_idx
    gap = 0
    for i in range(start_idx + 1, len(scores)):
        if scores[i] >= 1.0:
            end_idx = i
            gap = 0
        else:
            gap += 1
            if gap >= consecutive_break:
                break
    return (start_idx + 1, end_idx + 1)

# ---------- OLLAMA CALLS & PARSERS ----------
# def call_ollama(prompt: str, model: str = MODEL, timeout: int = 240) -> str:
#     try:
#         headers = {
#             "Authorization": f"Bearer {API_KEY}",   # required for OpenAI/OpenRouter
#             "Content-Type": "application/json"
#         }
#         payload = {
#             "model": model,
#             "messages": [{"role": "user", "content": prompt}],
#             "temperature": 0.3,
#             "max_tokens": 800
#         }
#         resp = requests.post(API_URL, json=payload, headers=headers, timeout=timeout)
#         resp.raise_for_status()
#         data = resp.json()
#         if "choices" in data:  # OpenAI/OpenRouter format
#             return data["choices"][0]["message"]["content"]
#         return data.get("response", "")
#     except Exception as e:
#         return f"API_ERROR: {e}"

async def call_ollama_text(prompt: str, model: str = MODEL, timeout: int = 240) -> str:
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": 800,
        "stream": False
    }

    async with httpx.AsyncClient(timeout=timeout) as client:
        resp = await client.post(API_URL, headers=headers, json=payload)
        resp.raise_for_status()
        data = resp.json()
        return data["choices"][0]["message"]["content"]


async def generate_mcqs_ollama(topic: str, context: str = "", full_text: str = "", model: str = MODEL):
    prompt = f"""
Generate 5 distinct multiple-choice questions for the topic below. For each question include:
- Exactly 4 labeled options A) B) C) D)
- A single-letter correct answer on its own line: Answer: <A/B/C/D>
- (Optional) Difficulty line: Difficulty: <1-5>

Use exactly this format; do not add extra commentary or code fences.

Q1. <question text>
A) <option A>
B) <option B>
C) <option C>
D) <option D>
Answer: <letter>
Difficulty: <1-5>

Topic: {topic}
Context: {context[:1500]}
"""
    out = await call_ollama_text(prompt, model=model)
    if out.startswith("OLLAMA_ERROR"):
        return [{"question": out, "options": [], "answer": "", "difficulty": ""}]
    mcqs = []
    blocks = re.split(r'\n(?=Q\d+\.)', out)
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = [ln.rstrip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue
        q_line_idx = 0
        for idx, ln in enumerate(lines):
            if re.match(r'^\s*Q\d+\.', ln, re.I):
                q_line_idx = idx
                break
            if not re.match(r'^[A-D][\)\.\-:]', ln, re.I) and not re.search(r'(here are|multiple[-\s]?choice|based on the topic)', ln, re.I):
                q_line_idx = idx
                break
        q_line = clean_text(lines[q_line_idx])
        opts = []
        opt_end_idx = q_line_idx
        for j in range(q_line_idx + 1, len(lines)):
            m = re.match(r'^([A-D])[\)\.\-:]\s*(.*)$', lines[j], re.I)
            if m:
                label = m.group(1).upper()
                text = m.group(2).strip()
                opts.append(f"{label}. {text}")
                opt_end_idx = j
            else:
                break
        answer = ""
        difficulty = ""
        look_start = opt_end_idx + 1
        look_end = min(len(lines), opt_end_idx + 6)
        for k in range(look_start, look_end):
            ln = lines[k]
            m_ans = re.search(r'(?i)\b(?:answer|correct)[:\s\-]*\(?\s*([A-D])\s*\)?', ln)
            if m_ans:
                answer = m_ans.group(1).upper()
                continue
            m_diff = re.search(r'(?i)\b(?:difficulty|level)[:\s\-]*\(?\s*([1-5])\s*\)?', ln)
            if m_diff:
                difficulty = m_diff.group(1)
                continue
            m_single = re.match(r'^\s*([A-D])[\)\.\s]*$', ln, re.I)
            if m_single and not answer:
                answer = m_single.group(1).upper()
        if not answer:
            m_any = re.search(r'(?i)\banswer[:\s\-]*\(?\s*([A-D])\s*\)?', block)
            if m_any:
                answer = m_any.group(1).upper()
        if q_line and len(opts) >= 2:
            mcqs.append({"question": q_line, "options": opts, "answer": answer, "difficulty": difficulty})
    return mcqs

async def generate_descriptive_with_answers(topic: str, context: str = "", model: str = MODEL, num_qs: int = 5):
    prompt = f"""
Generate {num_qs} descriptive / short-answer / essay-style questions for the topic below.
For each question, also provide:
- Correct answer
- Difficulty level (1-5)

Return exactly in this format:

Q1. <question text>
Answer: <answer text>
Difficulty: <1-5>

Do not add extra commentary.

Topic: {topic}
Context: {context[:1500]}
"""
    out = await call_ollama_text(prompt, model=model)
    if out.startswith("OLLAMA_ERROR"):
        return [{"question": out, "answer": "", "difficulty": ""}]
    blocks = re.split(r'\n(?=Q\d+\.)', out)
    results = []
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()
        question = ""
        answer = ""
        difficulty = ""
        for ln in lines:
            ln = ln.strip()
            if ln.lower().startswith("q"):
                question = re.sub(r'^q\d+\.\s*', '', ln, flags=re.I).strip()
            elif ln.lower().startswith("answer:"):
                answer = ln.split(":", 1)[1].strip()
            elif ln.lower().startswith("difficulty:"):
                difficulty = ln.split(":", 1)[1].strip()
        if question:
            results.append({"question": question, "answer": answer, "difficulty": difficulty})
    return results

def build_docx_bytes(questions_data: dict) -> bytes:
    doc = Document()
    doc.add_heading("Generated Questions", level=1)
    for topic_title, blocks in questions_data.items():
        doc.add_heading(topic_title, level=2)
        mcqs = blocks.get("mcqs", []) or []
        if mcqs:
            doc.add_paragraph("Multiple Choice Questions:")
            for idx, mcq in enumerate(mcqs, start=1):
                doc.add_paragraph(f"{idx}. {mcq.get('question','')}")
                for opt in mcq.get("options", []):
                    doc.add_paragraph(f"    {opt}")
                ans = mcq.get("answer", "")
                diff = mcq.get("difficulty", "N/A")
                if ans:
                    doc.add_paragraph(f"    Answer: {ans}    Difficulty: {diff}")
                else:
                    doc.add_paragraph(f"    Difficulty: {diff}")
                doc.add_paragraph("")
        descrs = blocks.get("descriptive", []) or []
        if descrs:
            doc.add_paragraph("Descriptive / Short-answer Questions:")
            for idx, dq in enumerate(descrs, start=1):
                if isinstance(dq, dict):
                    q = dq.get("question", "")
                    a = dq.get("answer", "")
                    diff = dq.get("difficulty", "N/A")
                else:
                    q = str(dq)
                    a, diff = "", "N/A"
                doc.add_paragraph(f"{idx}. {q}")
                if a:
                    doc.add_paragraph(f"    Answer: {a}")
                doc.add_paragraph(f"    Difficulty: {diff}")
                doc.add_paragraph("")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def build_dfs_from_questions(questions_data: dict):
    rows = []
    for topic_title, topic_data in questions_data.items():
        for mcq in topic_data.get("mcqs", []):
            opts = mcq.get("options") or []
            rows.append({
                "Topic": topic_title,
                "Type": "MCQ",
                "Question": mcq.get("question", ""),
                "Option A": opts[0] if len(opts) > 0 else "",
                "Option B": opts[1] if len(opts) > 1 else "",
                "Option C": opts[2] if len(opts) > 2 else "",
                "Option D": opts[3] if len(opts) > 3 else "",
                "Answer": mcq.get("answer", ""),
                "Difficulty": mcq.get("difficulty", "N/A"),
                "Descriptive Answer": ""
            })
        for dq in topic_data.get("descriptive", []):
            rows.append({
                "Topic": topic_title,
                "Type": "Descriptive",
                "Question": dq.get("question", ""),
                "Option A": "", "Option B": "", "Option C": "", "Option D": "",
                "Answer": "",
                "Difficulty": dq.get("difficulty", "N/A"),
                "Descriptive Answer": dq.get("answer", "")
            })
    return pd.DataFrame(rows)

def extract_audio_from_video(video_bytes: bytes) -> str:
    """Save video bytes to temp file and extract audio to WAV"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as vf:
        vf.write(video_bytes)
        vpath = vf.name
    wavpath = vpath.replace(".mp4", ".wav")
    try:
        # ffmpeg must be installed
        subprocess.run(
            ["ffmpeg", "-y", "-i", vpath, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", wavpath],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE
        )
    except Exception as e:
        print("FFmpeg error", e)
    return wavpath

try:
    WHISPER_MODEL = whisper.load_model("small")   # you can change "small" to "base" / "medium" / "large"
except Exception as e:
    print("❌ Failed to load Whisper model:", e)
    WHISPER_MODEL = None


def whisper_transcribe(audio_path: str, model_size: str = "small") -> str:
    if WHISPER_MODEL is None:
        return "Whisper model not available"
    try:
        result = WHISPER_MODEL.transcribe(audio_path)
        return result.get("text", "").strip()
    except Exception as e:
        print("❌ Whisper error:", e)
        return ""

# def whisper_transcribe(audio_path: str, model_size: str = "small") -> str:
#     """Use OpenAI Whisper to transcribe"""
#     try:
#         model = whisper.load_model(model_size)
#         result = model.transcribe(audio_path)
#         return result.get("text", "").strip()
#     except Exception as e:
#         print("❌ Whisper error:", e)
#         return ""

def save_questions_to_db(results: dict):
    db = SessionLocal()
    try:
        for topic, data in results.items():
            # Save MCQs
            for mcq in data.get("mcqs", []):
                opts = mcq.get("options", [])
                q = Question(
                    topic=topic,
                    type="MCQ",
                    question=mcq.get("question", ""),
                    option_a=opts[0] if len(opts) > 0 else None,
                    option_b=opts[1] if len(opts) > 1 else None,
                    option_c=opts[2] if len(opts) > 2 else None,
                    option_d=opts[3] if len(opts) > 3 else None,
                    answer=mcq.get("answer", ""),
                    difficulty=mcq.get("difficulty", "")
                )
                db.add(q)

            # Save Descriptive
            for dq in data.get("descriptive", []):
                q = Question(
                    topic=topic,
                    type="Descriptive",
                    question=dq.get("question", ""),
                    descriptive_answer=dq.get("answer", ""),
                    difficulty=dq.get("difficulty", "")
                )
                db.add(q)

        db.commit()
    except Exception as e:
        db.rollback()
        print("❌ DB error:", e)
    finally:
        db.close()

# ---------- ENDPOINTS: PDF / TOC / GENERATION ----------
@app.post("/extract_toc")
async def extract_toc(file: UploadFile = File(...)):
    pdf_bytes = await file.read()
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        # update page count state (not counting as upload until generation)
        STATE["last_pdf_pages"] = getattr(doc, "page_count", 0)
        # Try detect TOC pages and parse numeric headings
        try:
            start, end = detect_index_range(doc)
        except Exception:
            start, end = 1, min(6, doc.page_count)
        text = "\n".join([doc.load_page(p-1).get_text("text") or "" for p in range(start, end+1)])
        raw_matches = re.findall(r"(\d{1,2}\.\d+)\s+(.+?)\s+(\d{1,4})\b", text)
        matches = []
        if raw_matches:
            for num, title, pno in raw_matches:
                title_clean = re.sub(r"\.{2,}|\.{3,}", ".", title).strip(' .\t')
                title_clean = clean_text(title_clean)
                page_no = int(pno) if pno.isdigit() else None
                matches.append({"subnum": num.strip(), "title": title_clean, "page": page_no})
        else:
            # fallback: search simple lines
            for ln in text.splitlines():
                m = re.match(r'^\s*(\d{1,2}\.\d+)\s+(.+?)\s+(\d{1,4})\s*$', ln)
                if m:
                    matches.append({"subnum": m.group(1), "title": clean_text(m.group(2).strip()), "page": int(m.group(3))})
        # Build chapters map
        chapters = {}
        for m in matches:
            chap = int(m["subnum"].split(".")[0]) if m["subnum"].split(".")[0].isdigit() else 0
            chapters.setdefault(chap, []).append(m)
        return {"status": "success", "matches": matches, "chapters_count": len(chapters), "pages": STATE["last_pdf_pages"]}
    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.post("/generate_pdf_mcqs")
async def generate_pdf_mcqs(
    file: UploadFile = File(...),
    chapters: str = Form("[]"),
    question_type: str = Form("both"),
    #api_key: str = Depends(get_api_key)
):
    pdf_bytes = await file.read()
    selected_chapters = json.loads(chapters)
    qtype = (question_type or "both").lower()
    try:
        # update upload counter on new unique file
        md5 = hashlib.md5(pdf_bytes).hexdigest()
        if STATE.get("last_pdf_hash") != md5:
            STATE["pdf_uploads"] += 1
            STATE["last_pdf_hash"] = md5

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        STATE["last_pdf_pages"] = getattr(doc, "page_count", 0)
        full_text = "\n".join([doc.load_page(p).get_text("text") or "" for p in range(doc.page_count)])

        try:
            start, end = detect_index_range(doc)
            index_text = "\n".join([doc.load_page(p-1).get_text("text") or "" for p in range(start, end+1)])
        except Exception:
            index_text = full_text[:4000]

        raw_matches = re.findall(r"(\d{1,2}\.\d+)\s+(.+?)\s+(\d{1,4})\b", index_text)
        topics = []
        if raw_matches:
            for num, title, pno in raw_matches:
                title_clean = clean_text(re.sub(r"\.{2,}|\.{3,}", ".", title).strip(' .\t'))
                page_no = int(pno) if pno.isdigit() else None
                topics.append({"subnum": num, "title": title_clean, "page": page_no})
        else:
            for ln in index_text.splitlines():
                m = re.match(r'^\s*(\d{1,2}\.\d+)\s+(.+)$', ln)
                if m:
                    topics.append({"subnum": m.group(1), "title": clean_text(m.group(2).strip()), "page": None})

        # Filter by selected chapters if provided
        if selected_chapters:
            filtered = []
            for t in topics:
                chap_no = int(t["subnum"].split(".")[0]) if t["subnum"].split(".")[0].isdigit() else 0
                if chap_no in selected_chapters:
                    filtered.append(t)
            topics = filtered

        # Decide which types to produce
        produce_mcq = (qtype in ("mcq", "both"))
        produce_desc = (qtype in ("descriptive", "both"))

        # Generate questions for each topic (only requested types)
        results = {}
        for t in topics:
            title = t["title"]
            if t.get("page"):
                pg = t["page"]
                startp = max(0, pg-2)
                endp = min(doc.page_count, pg+1)
                context = "\n".join([doc.load_page(p).get_text("text") or "" for p in range(startp, endp)])
            else:
                context = index_text[:2000]

            entry = {}
            if produce_mcq:
                entry["mcqs"] = await generate_mcqs_ollama(title, context=context, full_text=full_text)
            else:
                entry["mcqs"] = []

            if produce_desc:
                entry["descriptive"] = await generate_descriptive_with_answers(title, context=context, num_qs=3)
            else:
                entry["descriptive"] = []

            # only include topic if at least one type requested (keeps structure clean)
            results[title] = entry

        # Build files and store them (only include selected types because results contains only requested parts)
        df_all = build_dfs_from_questions(results)

        # CSV
        csv_bytes = df_all.to_csv(index=False).encode("utf-8")
        csv_key = hashlib.md5(csv_bytes).hexdigest()
        store_result_bytes(csv_key, csv_bytes, "questions.csv", "text/csv")

        # Excel
        excel_buf = BytesIO()
        with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
            df_all.to_excel(writer, sheet_name="Questions", index=False)
        excel_buf.seek(0)
        excel_bytes = excel_buf.getvalue()
        excel_key = hashlib.md5(excel_bytes).hexdigest()
        store_result_bytes(excel_key, excel_bytes, "questions.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # DOCX
        docx_bytes = build_docx_bytes(results)
        docx_key = hashlib.md5(docx_bytes).hexdigest()
        store_result_bytes(docx_key, docx_bytes, "questions.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # compute counts for this generation (returned to client and used by front-end to update dashboard)
        mcq_count_now = sum(len(v.get("mcqs", [])) for v in results.values())
        desc_count_now = sum(len(v.get("descriptive", [])) for v in results.values())

        # NOTE: STATE counters remain cumulative if you want; frontend will use the per-request counts to update the dashboard
        STATE["mcq_count"] = STATE.get("mcq_count", 0) + mcq_count_now
        STATE["desc_count"] = STATE.get("desc_count", 0) + desc_count_now

        return {
            "status": "success",
            "results_count_topics": len(results),
            "mcqCount": mcq_count_now,
            "descCount": desc_count_now,
            "download_keys": {"csv": csv_key, "excel": excel_key, "docx": docx_key},
            "pages": STATE["last_pdf_pages"],
            "global_state": {
                "pdf_uploads": STATE["pdf_uploads"],
                "last_pdf_pages": STATE["last_pdf_pages"],
                "mcq_count": STATE["mcq_count"],
                "desc_count": STATE["desc_count"]
            },
            "results": results   # for immediate front-end rendering
        }

    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.post("/transcribe_video")
async def transcribe_video(file: UploadFile = File(...), whisper_model: str = Form("small")):
    try:
        video_bytes = await file.read()
        audio_path = extract_audio_from_video(video_bytes)
        transcript = whisper_transcribe(audio_path, whisper_model)
        return {
            "status": "success",
            "transcript": transcript,
            "summary": transcript[:300]  # just preview first 300 chars
        }
    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.post("/generate_video_mcqs")
async def generate_video_mcqs(
    file: UploadFile = File(...),
    question_type: str = Form("both"),
    num_qs: int = Form(5),
    whisper_model: str = Form("small"),
    api_key: str = Depends(get_api_key)
):
    try:
        # 1️⃣ Save video & extract audio
        video_bytes = await file.read()
        audio_path = extract_audio_from_video(video_bytes)

        # 2️⃣ Transcribe with Whisper
        transcript_text = whisper_transcribe(audio_path, whisper_model)
        if not transcript_text:
            return {"status": "error", "error": "Transcription failed"}

        # 3️⃣ Generate MCQs + Descriptive
        results = {
            "Video Topic": {
                "mcqs": await generate_mcqs_ollama("Video Topic", context=transcript_text[:2000]),
                "descriptive": await generate_descriptive_with_answers(
                    "Video Topic", context=transcript_text[:2000], num_qs=num_qs
                ),
            }
        }

        # 4️⃣ Save to MySQL
        save_questions_to_db(results)

        # 5️⃣ Export
        df_all = build_dfs_from_questions(results)
        csv_bytes = df_all.to_csv(index=False).encode("utf-8")
        csv_key = hashlib.md5(csv_bytes).hexdigest()
        store_result_bytes(csv_key, csv_bytes, "video_questions.csv", "text/csv")

        docx_bytes = build_docx_bytes(results)
        docx_key = hashlib.md5(docx_bytes).hexdigest()
        store_result_bytes(
            docx_key,
            docx_bytes,
            "video_questions.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "status": "success",
            "results_count_topics": len(results),
            "download_keys": {"csv": csv_key, "docx": docx_key},
            "results": results,
            "mcqCount": len(results["Video Topic"]["mcqs"]),
            "descCount": len(results["Video Topic"]["descriptive"]),
        }

    except Exception as e:
        return {"status": "error", "error": str(e)}


@app.get("/questions")
def get_questions(topic: str = None, qtype: str = None):
    db = SessionLocal()
    try:
        query = db.query(Question)   # 👈 start with a base query

        if topic:
            query = query.filter(Question.topic == topic)
        if qtype:
            query = query.filter(Question.type == qtype)

        results = query.all()
        return JSONResponse(content=jsonable_encoder(results))  # serialize properly
    finally:
        db.close()

@app.get("/all_questions")
def all_questions():
    db = SessionLocal()
    try:
        questions = db.query(Question).order_by(Question.created_at.desc()).all()
        return JSONResponse(content=jsonable_encoder(questions))
    finally:
        db.close()

@app.post("/export_pdf_to_db")
async def export_pdf_to_db(results: str = Form(...), api_key: str = Depends(get_api_key)):
    try:
        data = json.loads(results)
        save_questions_to_db(data)
        return {"status": "success", "message": "PDF questions exported to database."}
    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.post("/export_video_to_db")
async def export_video_to_db(results: str = Form(...), api_key: str = Depends(get_api_key)):
    try:
        data = json.loads(results)
        save_questions_to_db(data)
        return {"status": "success", "message": "Video questions exported to database."}
    except Exception as e:
        return {"status": "error", "error": str(e)}
    
@app.post("/chat_stream")
async def chat_stream(prompt: str = Form(...), api_key: str = Depends(get_api_key)):
    return StreamingResponse(call_ollama_text(prompt), media_type="text/event-stream")


# ---------- RUN ----------
if __name__ == "__main__":
    print(f"Starting backend at http://{HOST}:{PORT}  (static files dir: {static_dir})")
    uvicorn.run("backend1:app", host=HOST, port=PORT, reload=True)